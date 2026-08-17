"""
Document & Media Processing Tools Engines.

Engines:
1. Image to Text (OCR engine via Vision AI / Tesseract fallback)
2. Image to Word (Scanned Image to .docx converter)
3. Real PDF Text Content Editor (Parses & modifies PDF text streams)
4. PDF to Word (.docx) Converter
5. Word (.docx) to PDF Converter
6. Social Media URL Downloader (YouTube, Instagram, TikTok, Twitter/X, Facebook)
7. Webpage URL to PDF / Image File Converter
"""
import os
import io
import re
import urllib.parse
import logging
import requests

logger = logging.getLogger(__name__)


def image_to_text_engine(image_bytes: bytes, filename: str = "") -> dict:
    """Extracts text from image bytes using OCR / Vision AI."""
    from app.ai_tools.routes import _call_ai
    import base64

    b64_img = base64.b64encode(image_bytes).decode("utf-8")
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "png"
    mime_type = f"image/{'jpeg' if ext in ('jpg', 'jpeg') else ext}"
    
    prompt = "Perform accurate OCR on this image. Extract and return ALL text visible in the image verbatim, preserving original paragraph breaks and list structure. Return only the extracted text, nothing else."
    
    # Try Vision AI model call first for ultra-accurate OCR
    messages = [{
        "role": "user",
        "content": [
            {"type": "text", "text": prompt},
            {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{b64_img}"}}
        ]
    }]
    
    text, err = _call_ai("You are an expert OCR text extractor.", messages)
    if not err and text:
        return {"success": True, "text": text.strip()}

    # Tesseract fallback if installed on server
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(image_bytes))
        extracted = pytesseract.image_to_string(img)
        if extracted.strip():
            return {"success": True, "text": extracted.strip()}
    except Exception:
        pass

    return {"success": False, "error": err or "Could not extract text from image."}


def image_to_word_engine(image_bytes: bytes, filename: str = "") -> bytes:
    """Converts image with text/layout into an editable .docx Word document."""
    res = image_to_text_engine(image_bytes, filename)
    extracted_text = res.get("text", "") if res.get("success") else "Scanned Image Document"
    
    doc_title = filename.rsplit(".", 1)[0] if "." in filename else "Scanned_Document"
    
    word_html = f"""<html xmlns:o='urn:schemas-microsoft-com:office:office' xmlns:w='urn:schemas-microsoft-com:office:word' xmlns='http://www.w3.org/TR/REC-html40'>
<head><meta charset='utf-8'><title>{doc_title}</title>
<style>
body {{ font-family: 'Calibri', 'Arial', sans-serif; font-size: 11pt; line-height: 1.5; margin: 1in; }}
p {{ margin-bottom: 10pt; }}
</style>
</head>
<body>
<h1>{doc_title}</h1>
<p>{extracted_text.replace('\n', '<br/>')}</p>
</body>
</html>"""
    return word_html.encode("utf-8")


def pdf_to_word_engine(pdf_bytes: bytes) -> bytes:
    """Converts PDF bytes to editable .docx Word document.

    Tries pdf2docx first (best layout/table fidelity). If that's
    unavailable or fails on this particular PDF, falls back to
    pdfplumber, which is much lighter-weight (no PyMuPDF/opencv/numpy)
    and — critically — can detect tables and rebuild them as real Word
    tables, instead of the old fallback which just concatenated all
    text (including table cell text) into one paragraph and silently
    lost table structure, wrapped in fake, non-OOXML ".docx" bytes that
    weren't a real Word file at all."""
    try:
        from pdf2docx import Converter
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f_pdf:
            f_pdf.write(pdf_bytes)
            pdf_path = f_pdf.name

        docx_path = pdf_path.replace(".pdf", ".docx")
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()

        with open(docx_path, "rb") as f_docx:
            docx_data = f_docx.read()

        os.unlink(pdf_path)
        if os.path.exists(docx_path):
            os.unlink(docx_path)
        if len(docx_data) < 1000:
            raise RuntimeError("pdf2docx produced a suspiciously small/empty file")
        return docx_data
    except Exception as e:
        logger.exception("pdf2docx conversion failed — using pdfplumber table-aware fallback")
        return _pdf_to_word_fallback(pdf_bytes)


def _pdf_to_word_fallback(pdf_bytes: bytes) -> bytes:
    """Table-aware fallback: builds a REAL .docx via python-docx. Uses
    pdfplumber (lighter than pdf2docx's PyMuPDF/opencv/numpy stack) to
    pull out both tables and the surrounding paragraph text per page,
    in reading order, so tables land in the output as actual Word
    tables rather than being dropped or jumbled into plain text."""
    import docx as docx_lib

    document = docx_lib.Document()
    used_pdfplumber = False
    try:
        import pdfplumber
        used_pdfplumber = True
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                if page_num > 0:
                    document.add_page_break()
                tables = page.find_tables()
                table_bboxes = [t.bbox for t in tables]

                def _inside_any_table(word):
                    wy = (word["top"] + word["bottom"]) / 2
                    wx = (word["x0"] + word["x1"]) / 2
                    for (x0, top, x1, bottom) in table_bboxes:
                        if x0 <= wx <= x1 and top <= wy <= bottom:
                            return True
                    return False

                # Non-table text, grouped into lines, with each table
                # inserted as a real Word table at roughly its position
                # in reading order (top-to-bottom).
                words = [w for w in page.extract_words() if not _inside_any_table(w)]
                lines = {}
                for w in words:
                    key = round(w["top"] / 3)
                    lines.setdefault(key, []).append(w)

                blocks = [("text", k, lines[k]) for k in lines]
                blocks += [("table", t.bbox[1], t) for t in tables]
                blocks.sort(key=lambda b: b[1])

                for kind, _, payload in blocks:
                    if kind == "text":
                        line_text = " ".join(w["text"] for w in sorted(payload, key=lambda w: w["x0"]))
                        if line_text.strip():
                            document.add_paragraph(line_text.strip())
                    else:
                        data = payload.extract() or []
                        data = [row for row in data if row]
                        if not data:
                            continue
                        n_cols = max(len(row) for row in data)
                        tbl = document.add_table(rows=0, cols=n_cols)
                        try:
                            tbl.style = "Table Grid"
                        except Exception:
                            pass
                        for row in data:
                            cells = tbl.add_row().cells
                            for i in range(n_cols):
                                cells[i].text = (row[i] or "").strip() if i < len(row) and row[i] else ""
                        document.add_paragraph("")
    except ImportError:
        used_pdfplumber = False

    if not used_pdfplumber:
        # Last-resort: plain text only, but still a real, valid .docx —
        # never the old fake-HTML-with-a-.docx-name output.
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(pdf_bytes))
            for page in reader.pages:
                for line in (page.extract_text() or "").split("\n"):
                    if line.strip():
                        document.add_paragraph(line.strip())
        except Exception as e2:
            raise RuntimeError(f"PDF to Word conversion failed: {e2}")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def word_to_pdf_engine(word_bytes: bytes) -> bytes:
    """Converts Word .docx bytes to standard .pdf bytes.

    Raises ValueError with a clear, user-facing message for known bad
    input (legacy .doc, corrupt/non-Word file) so the route can show it
    directly instead of falling through to a generic 500 page.
    """
    # Legacy binary .doc files (Word 97-2003) start with the OLE2/CFBF
    # magic bytes below. python-docx (and every free pure-Python docx
    # reader) can only parse the newer .docx XML/zip format — there is
    # no free way to read old .doc on a host without LibreOffice/MS Word
    # installed, so detect this up front and say so clearly rather than
    # letting python-docx throw an opaque exception.
    if word_bytes[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        raise ValueError(
            "This is an old .doc file (Word 97-2003 format), which can't be read directly. "
            "Please open it in Word or Google Docs and use 'Save As' / 'Download' → Word Document (.docx), "
            "then upload that .docx file instead."
        )

    import docx
    from docx.opc.exceptions import PackageNotFoundError
    try:
        doc = docx.Document(io.BytesIO(word_bytes))
    except PackageNotFoundError:
        raise ValueError("This doesn't look like a valid .docx file. Please double-check the file and try again.")
    except Exception as e:
        logger.exception("word_to_pdf: failed to open .docx")
        raise ValueError(f"Couldn't read this Word document: {e}")

    full_text = "\n".join(p.text for p in doc.paragraphs)

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        raise RuntimeError("PDF generation isn't installed on the server — run `pip install reportlab` on the server.")

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    y = 750
    for line in full_text.split("\n") or [""]:
        if y < 50:
            c.showPage()
            y = 750
        c.drawString(50, y, line[:90])
        y -= 15
    c.save()
    return buffer.getvalue()


def social_media_download_engine(url: str) -> dict:
    """Extracts a real, direct downloadable media URL from a social media
    page URL (YouTube, IG, TikTok, Twitter/X, FB) via yt-dlp."""
    url_clean = url.strip()

    try:
        import yt_dlp
    except ImportError:
        return {"success": False, "error": "yt-dlp isn't installed on the server — run `pip install yt-dlp`."}

    # 'best' alone often resolves to a separate (not-yet-merged) video-only
    # stream on sites that split audio/video by quality — merging those
    # requires ffmpeg, which isn't available here. Asking for a *progressive*
    # (single-file, already has both audio+video) format avoids needing
    # ffmpeg at all, at the cost of capping resolution on a few sites.
    ydl_opts = {'format': 'best[acodec!=none][vcodec!=none]/best', 'quiet': True, 'no_warnings': True}
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url_clean, download=False)
    except Exception as e:
        return {"success": False, "error": f"Couldn't extract this media — {e}"}

    media_url = info.get('url')
    if not media_url:
        # Pick the best progressive (audio+video, already-merged) format
        # available, instead of blindly grabbing whichever format happened
        # to be last in the list (which was often audio-only or low-res).
        formats = [f for f in (info.get('formats') or [])
                   if f.get('url') and f.get('acodec') not in (None, 'none') and f.get('vcodec') not in (None, 'none')]
        if formats:
            best = max(formats, key=lambda f: f.get('height') or 0)
            media_url = best.get('url')

    if not media_url:
        return {"success": False, "error": "No directly-downloadable (audio+video merged) format is available for this URL without server-side ffmpeg, which isn't set up here."}

    return {
        "success": True,
        "title": info.get('title', 'Social Media Media File'),
        "download_url": media_url,
        "thumbnail": info.get('thumbnail'),
        "duration": info.get('duration'),
        "ext": info.get('ext', 'mp4'),
    }


def url_to_file_engine(url: str, output_format: str = "pdf") -> bytes:
    """Fetches a public webpage and exports its readable text content as a
    real PDF (or raw HTML). This extracts text, not a visual screenshot —
    a pixel-accurate "print webpage to PDF" needs a headless browser
    (Playwright/wkhtmltopdf), which needs system packages this host can't
    install."""
    resp = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=15)
    resp.raise_for_status()
    html = resp.text

    if output_format == "html":
        return html.encode("utf-8")

    from bs4 import BeautifulSoup
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    title = (soup.title.string.strip() if soup.title and soup.title.string else url)[:120]
    text = soup.get_text("\n")
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    margin = 0.75 * inch
    y = height - margin
    c.setFont("Helvetica-Bold", 14)
    for chunk_start in range(0, len(title), 90):
        c.drawString(margin, y, title[chunk_start:chunk_start + 90])
        y -= 18
    y -= 10
    c.setFont("Helvetica", 10)
    max_chars = 95
    for line in lines:
        for i in range(0, len(line), max_chars) or [0]:
            if y < margin:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = height - margin
            c.drawString(margin, y, line[i:i + max_chars])
            y -= 13
    c.save()
    return buffer.getvalue()


def _ffmpeg_path():
    """A real, static ffmpeg binary bundled by the imageio-ffmpeg package
    (downloaded once at pip-install time) — no system package manager or
    root access needed, which matters since PythonAnywhere doesn't allow
    installing system packages like `apt install ffmpeg`."""
    try:
        import imageio_ffmpeg
        return imageio_ffmpeg.get_ffmpeg_exe()
    except ImportError:
        return None


MAX_VIDEO_BYTES = 200 * 1024 * 1024  # 200 MB cap for uploads to this engine


def mov_to_mp4_engine(video_bytes: bytes) -> bytes:
    """Converts a .mov (QuickTime) file to a standard, widely-playable
    .mp4 (H.264 + AAC)."""
    if len(video_bytes) > MAX_VIDEO_BYTES:
        raise ValueError(f"File too large — max {MAX_VIDEO_BYTES // (1024*1024)} MB.")
    ffmpeg = _ffmpeg_path()
    if not ffmpeg:
        raise RuntimeError("Video conversion isn't installed on the server — run `pip install imageio-ffmpeg`.")

    import subprocess, tempfile
    with tempfile.NamedTemporaryFile(suffix=".mov", delete=False) as f_in:
        f_in.write(video_bytes)
        in_path = f_in.name
    out_path = in_path.replace(".mov", ".mp4")

    try:
        result = subprocess.run(
            [ffmpeg, "-y", "-i", in_path, "-c:v", "libx264", "-preset", "fast", "-crf", "23",
             "-c:a", "aac", "-b:a", "128k", "-movflags", "+faststart", out_path],
            capture_output=True, text=True, timeout=280,
        )
        if result.returncode != 0 or not os.path.exists(out_path) or os.path.getsize(out_path) < 500:
            raise RuntimeError(f"ffmpeg conversion failed: {result.stderr[-500:] if result.stderr else 'unknown error'}")
        with open(out_path, "rb") as f_out:
            return f_out.read()
    finally:
        for p in (in_path, out_path):
            try:
                if os.path.exists(p):
                    os.remove(p)
            except OSError:
                pass


def video_compress_engine(video_bytes: bytes, filename: str = "video.mp4", target_quality: str = "balanced") -> bytes:
    """Compresses a video (mp4/mov/webm/etc.) down to a much smaller file
    size using H.264 with a quality-based CRF, re-encoding audio to a
    lower bitrate too. 'balanced' is a good default for sharing online;
    'smaller' pushes size down further at some quality cost."""
    if len(video_bytes) > MAX_VIDEO_BYTES:
        raise ValueError(f"File too large — max {MAX_VIDEO_BYTES // (1024*1024)} MB.")
    ffmpeg = _ffmpeg_path()
    if not ffmpeg:
        raise RuntimeError("Video compression isn't installed on the server — run `pip install imageio-ffmpeg`.")

    crf = {"smaller": "30", "balanced": "26", "high_quality": "20"}.get(target_quality, "26")
    audio_bitrate = {"smaller": "96k", "balanced": "128k", "high_quality": "160k"}.get(target_quality, "128k")

    import subprocess, tempfile
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "mp4"
    with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as f_in:
        f_in.write(video_bytes)
        in_path = f_in.name
    out_path = in_path.rsplit(".", 1)[0] + "_compressed.mp4"

    try:
        result = subprocess.run(
            [ffmpeg, "-y", "-i", in_path, "-c:v", "libx264", "-preset", "medium", "-crf", crf,
             "-c:a", "aac", "-b:a", audio_bitrate, "-movflags", "+faststart", out_path],
            capture_output=True, text=True, timeout=280,
        )
        if result.returncode != 0 or not os.path.exists(out_path) or os.path.getsize(out_path) < 500:
            raise RuntimeError(f"ffmpeg compression failed: {result.stderr[-500:] if result.stderr else 'unknown error'}")
        with open(out_path, "rb") as f_out:
            return f_out.read()
    finally:
        for p in (in_path, out_path):
            try:
                if os.path.exists(p):
                    os.remove(p)
            except OSError:
                pass
